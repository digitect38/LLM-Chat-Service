#!/usr/bin/env python3
"""
Add v4.86.0 changelog entry to Fab_Copilot_v4_84.docx
- Adds revision history row
- Adds new section 9.14 before PART V
"""
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT  = r'C:\Develop25\VisualFactoryHome\LLM-Chat-Service\Plan\Fab_Copilot_v4_84.docx'
OUTPUT = r'C:\Develop25\VisualFactoryHome\LLM-Chat-Service\Plan\Fab_Copilot_v4_86.docx'

doc = Document(INPUT)

# ─── 1. Add revision history row (Table 0) ───────────────────────

table = doc.tables[0]
new_row = table.add_row()
cells = new_row.cells
cells[0].text = 'v4.86.0'
cells[1].text = '2025-03'
cells[2].text = (
    'Streaming Token Ordering & TTS Init Race Fix (Section 9.14 신설):\n'
    '• HandleChunkReceived 재구조화 — 토큰 append를 모든 await 이전에 동기 실행하여 '
    'Blazor InvokeAsync 콜백 interleaving에 의한 토큰 순서 역전 방지\n'
    '• _ttsInitPending 동기 guard 추가 — TTS playStream init 중 중복 진입 방지, '
    'init 완료 후 축적된 전체 텍스트를 한 번에 feedText 전달\n'
    '• JIT 렌더링 간격 300ms → 150ms 단축 (초당 ~3→~6회 갱신, 체감 속도 개선)\n'
    '• IsComplete 시 텍스트 무결성 검증 로그 추가 (console.warn 기반 디버깅 지원)'
)
cells[3].text = ''

for cell in cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(9)

print(f'[OK] Revision history row added (now {len(table.rows)} rows)')

# ─── 2. Find insertion point: before "PART V" ────────────────────

insert_before_elem = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'PART V — Common Infrastructure':
        insert_before_elem = p._element
        print(f'[OK] Insertion point: paragraph [{i}] = "{p.text[:50]}"')
        break

if insert_before_elem is None:
    print('[ERROR] Could not find PART V')
    sys.exit(1)

body = doc.element.body

# ─── 3. Helper: insert paragraph before element ──────────────────

def add_para(text, style_name):
    """Create a new paragraph and insert it before the insertion point."""
    new_p = OxmlElement('w:p')
    body.insert(list(body).index(insert_before_elem), new_p)

    # Find the paragraph object
    for p in doc.paragraphs:
        if p._element is new_p:
            p.style = style_name
            if text:
                run = p.add_run(text)
                run.font.size = Pt(10)
            return p
    return None


# ─── 4. Insert Section 9.14 content ──────────────────────────────

add_para('', 'Normal')

add_para('9.14 Streaming Token Ordering & Rendering Architecture', 'Heading 2')

add_para(
    '채팅 응답의 실시간 스트리밍에서 토큰 순서 보장, 렌더링 파이프라인, '
    'TTS 스트리밍 초기화의 동시성 제어를 정의한다. '
    'v4.86.0에서 발견된 토큰 순서 역전 버그의 근본 원인 분석과 해결 방안을 포함한다.',
    'Normal'
)

add_para('', 'Normal')

add_para('9.14.1 문제 현상 및 근본 원인', 'Heading 3')

add_para(
    '현상: 채팅 응답 텍스트가 깨져 보이는(garbled) 증상. '
    '토큰이 뒤섞여 "## 요약"이 텍스트 중간에 나타나거나, 응답 앞부분이 잘려 나옴. '
    '중간에 정지(stop)하지 않아도 첫 질문부터 발생. '
    'Redis에 저장된 대화 기록은 정상 — 새로고침 후 기록 열면 정상 텍스트 확인됨.',
    'Normal'
)

add_para('', 'Normal')

add_para(
    '근본 원인: Blazor Server의 InvokeAsync 콜백 interleaving. '
    'HandleChunkReceived에서 각 chunk에 대해 InvokeAsync(async () => { ... })를 호출하는데, '
    '콜백 내부에 await가 있으면 Blazor 동기화 컨텍스트가 해당 콜백을 일시 중단하고 '
    '다음 대기 중인 콜백을 실행한다. 기존 코드에서는 TTS 초기화(await fabTts.playStream)가 '
    '토큰 append 전에 실행되어, await로 양보하는 동안 후속 chunk 콜백이 먼저 토큰을 append하면서 '
    '순서가 뒤바뀌었다.',
    'Normal'
)

add_para('', 'Normal')

add_para(
    '검증: 3개 Python 테스트 스크립트(test_duplicate_detect.py, test_interleave_detect.py, '
    'test_stop_resume.py)로 백엔드(ChatGateway → NATS → LLM Worker → NATS → WebSocket) '
    '토큰 전달 정합성을 확인. 25/25 라운드 전부 통과. '
    '문제는 백엔드가 아닌 Blazor UI 렌더링 레이어에서 발생.',
    'Normal'
)

add_para('', 'Normal')

add_para('9.14.2 토큰 순서 보장 원칙', 'Heading 3')

add_para(
    'HandleChunkReceived의 InvokeAsync 콜백 내부에서, '
    '모든 상태 변경(토큰 append, IsComplete 처리, Error 처리)은 '
    '어떠한 await보다 먼저 동기적으로 실행되어야 한다. '
    'TTS 피드, 콘솔 로깅 등 비동기 작업은 토큰이 msg.Text에 안전하게 저장된 후 실행한다.',
    'Normal'
)

add_para('', 'Normal')

add_para('처리 순서 (수정 후):', 'Normal')

for step in [
    '1. Guard 검사: _isStopping, _isSending, ConversationId 일치 여부 (동기)',
    '2. assistantMsg 탐색 또는 생성 (동기)',
    '3. 토큰 append: assistantMsg.Text += chunk.Token (동기 — 순서 보장의 핵심)',
    '4. 렌더 스케줄링: ScheduleRenderUpdate(), ResetStaleStreamTimer() (동기)',
    '5. TTS 초기화 및 피드 (비동기 — await 포함, 토큰은 이미 저장됨)',
]:
    add_para(step, 'List Paragraph')

add_para('', 'Normal')

add_para(
    '기존 코드에서는 단계 5(TTS 초기화)가 단계 3(토큰 append) 이전에 실행되어 문제가 발생했다. '
    'TTS가 비활성화된 경우에도 진단용 console.log await가 hot path에 존재하여 '
    '동일한 interleaving 위험이 있었으므로, 해당 await도 제거하였다.',
    'Normal'
)

add_para('', 'Normal')

add_para('9.14.3 TTS 스트리밍 초기화 동시성 제어', 'Heading 3')

add_para(
    'TTS 초기화(fabTts.playStream)는 JS interop await를 포함하므로, '
    '첫 번째 chunk의 init이 완료되기 전에 후속 chunk들이 도착할 수 있다. '
    '이 동시성을 제어하기 위해 _ttsInitPending 플래그를 추가하여 3가지 상태를 관리한다:',
    'Normal'
)

add_para('', 'Normal')

for state in [
    '(a) !_ttsStreamingActive && !_ttsInitPending: 최초 init 시작 경로. '
    '_ttsInitPending = true를 동기적으로 설정한 후 await playStream 실행. '
    'init 완료 후 feedText(assistantMsg.Text)로 그때까지 축적된 전체 텍스트를 한 번에 전달.',

    '(b) _ttsStreamingActive: 정상 경로 (init 완료 후). '
    '개별 토큰을 feedText(chunk.Token)으로 전달.',

    '(c) _ttsInitPending && !_ttsStreamingActive: init 대기 중 경로. '
    'feedText 호출을 건너뜀. init 완료 후 feedText(assistantMsg.Text)에 의해 일괄 전달됨.',
]:
    add_para(state, 'List Paragraph')

add_para('', 'Normal')

add_para(
    '_ttsInitPending 플래그는 await 이전에 동기적으로 설정되므로, '
    'Blazor 동기화 컨텍스트에서 후속 콜백이 실행될 때 즉시 확인 가능하다. '
    '이를 통해 playStream의 중복 호출과 feedText의 텍스트 반복 전달(첫 문장 무한 반복 TTS)을 방지한다.',
    'Normal'
)

add_para('', 'Normal')

add_para('9.14.4 JIT 렌더링 파이프라인', 'Heading 3')

add_para('JIT(Just-In-Time) 렌더 모드에서의 스트리밍 렌더링 파이프라인:', 'Normal')

add_para('', 'Normal')

for step in [
    'ScheduleRenderUpdate: System.Threading.Timer 기반 디바운스. '
    '토큰 도착 시 타이머를 리셋하여 일정 간격(150ms)마다 렌더링 실행.',

    '타이머 콜백: InvokeAsync로 Blazor 동기화 컨텍스트에 마샬링. '
    'msg.Text를 JS interop으로 chatRenderer.renderMarkdownLite에 전달. '
    '반환된 HTML을 msg.RenderedHtml에 저장 후 StateHasChanged() 호출.',

    'renderMarkdownLite (JS): marked.js로 마크다운 파싱, hljs로 코드 하이라이팅, '
    '수학식은 KaTeX 대신 placeholder로 보호 (스트리밍 성능 최적화). '
    'IsComplete 시 전체 렌더러(renderMarkdown)로 KaTeX 포함 최종 렌더링 수행.',

    '렌더 간격: v4.86.0에서 JIT 모드 300ms → 150ms로 단축. '
    '초당 화면 갱신 빈도 약 3회 → 약 6회. Fade 모드는 기존 100ms 유지.',
]:
    add_para(step, 'List Paragraph')

add_para('', 'Normal')

add_para('9.14.5 텍스트 무결성 검증', 'Heading 3')

add_para(
    'IsComplete 수신 시 브라우저 콘솔(console.warn)에 축적된 텍스트의 길이와 '
    '처음 200자를 출력한다. 향후 텍스트 깨짐 현상 재발 시 F12 개발자 도구에서 '
    '실제 축적된 텍스트를 확인하여 토큰 순서 문제(데이터 레이어)와 '
    '마크다운 렌더링 문제(표시 레이어)를 구분할 수 있다.',
    'Normal'
)

add_para('', 'Normal')

add_para(
    '출력 형식: [Blazor] COMPLETE chunks={청크수} len={텍스트길이} text={처음200자}',
    'Normal'
)

add_para('', 'Normal')

add_para('9.14.6 변경 파일 요약', 'Heading 3')

for item in [
    'Index.razor — HandleChunkReceived: 토큰 append를 await 이전으로 이동, TTS init을 이후로 이동',
    'Index.razor — _ttsInitPending 필드 추가: TTS init 중복 진입 방지용 동기 guard',
    'Index.razor — ScheduleRenderUpdate: JIT 렌더 간격 300ms → 150ms',
    'Index.razor — IsComplete 분기: 텍스트 무결성 검증 console.warn 로그 추가',
    'Index.razor — AppVersion: v4.85.2 → v4.86.0',
]:
    add_para(item, 'List Paragraph')

add_para('', 'Normal')

# ─── 5. Save ──────────────────────────────────────────────────────

doc.save(OUTPUT)
print(f'\n[OK] Saved to: {OUTPUT}')
print(f'[OK] Total paragraphs: {len(doc.paragraphs)}')
